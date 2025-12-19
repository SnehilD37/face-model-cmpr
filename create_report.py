import pandas as pd
import os
import re
from docx import Document
from docx.shared import Pt

# 1. Define all 5 models to look for
FILES = {
    "FaceNet": "results_facenet.csv",
    "ArcFace": "results_arcface.csv",
    "InsightFace": "results_insightface.csv",
    "CosFace": "results_cosface.csv",
    "SphereFace": "results_sphereface.csv"
}

def clean_confidence(conf_val):
    """
    Parses "98.5%" -> 98.5 (float)
    Parses "Snehil:99%" -> 99.0 (float) - legacy support
    Returns 0.0 if empty/NaN.
    """
    if pd.isna(conf_val) or conf_val == "":
        return 0.0
    
    # Convert to string
    s = str(conf_val)
    
    # Extract the first valid number followed by a % symbol
    matches = re.findall(r"(\d+\.?\d*)%", s)
    if matches:
        # If multiple numbers (legacy format), average them
        scores = [float(m) for m in matches]
        return sum(scores) / len(scores)
        
    return 0.0

def generate_report():
    print("Generating Final Benchmark Report...")
    doc = Document()
    doc.add_heading('Final Face Recognition Benchmark', 0)
    
    doc.add_paragraph("Metric Definitions:")
    doc.add_paragraph("- Recall: % of images where the correct person was found.")
    doc.add_paragraph("- Avg Conf: The average certainty score for identified faces.")
    doc.add_paragraph("- Time: Average processing time per image (seconds).")
    
    # Create Table with 4 Columns
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Set Header Row
    hdr = table.rows[0].cells
    hdr[0].text = 'Model'
    hdr[1].text = 'Recall (Acc)'
    hdr[2].text = 'Avg Conf'
    hdr[3].text = 'Time (s)'
    
    for model, csv_file in FILES.items():
        row = table.add_row().cells
        row[0].text = model
        
        if os.path.exists(csv_file):
            try:
                df = pd.read_csv(csv_file)
                
                # --- IDENTIFY IMAGE TYPES ---
                # Exclude 'negative' or 'empty' images from Recall calculation
                is_neg = df['image'].str.lower().str.contains('negative') | df['image'].str.lower().str.contains('empty')
                pos_df = df[~is_neg]
                
                # --- CALCULATE METRICS ---
                recall_score = 0.0
                avg_conf_score = 0.0
                
                if len(pos_df) > 0:
                    # 1. Recall: Count rows where 'found_people' is not empty
                    hits_mask = pos_df['found_people'].notna() & (pos_df['found_people'] != "")
                    pos_hits = hits_mask.sum()
                    recall_score = round((pos_hits / len(pos_df) * 100), 1)
                    
                    # 2. Avg Confidence: Check for 'avg_confidence' OR 'confidence_levels'
                    # We prefer 'avg_confidence' (new format)
                    target_col = 'avg_confidence' if 'avg_confidence' in pos_df.columns else 'confidence_levels'
                    
                    if target_col in pos_df.columns:
                        # Apply cleaning function to the column
                        confs = pos_df[target_col].apply(clean_confidence)
                        # Filter out 0.0s (misses/unknowns) so they don't drag down the average
                        valid_confs = confs[confs > 0]
                        if len(valid_confs) > 0:
                            avg_conf_score = round(valid_confs.mean(), 1)

                # 3. Time: Average of the 'time' column
                avg_time = round(df['time'].mean(), 3)
                
                # --- FILL TABLE ---
                row[1].text = f"{recall_score}%"
                row[2].text = f"{avg_conf_score}%"
                row[3].text = str(avg_time)
                
            except Exception as e:
                row[1].text = "Error"
                print(f"Error processing {model}: {e}")
        else:
            row[1].text = "-" # File not found

    doc.save('Final_Benchmark_Report.docx')
    print("Success! Report saved as 'Final_Benchmark_Report.docx'")

if __name__ == "__main__":
    generate_report()